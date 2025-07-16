from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import asyncio
import os
import sys
import json
import threading
import logging
import subprocess
import traceback
from typing import Dict, List, Set
from requirement_analyzer import RequirementsAnalyzer
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")

# Ensure static directory exists
os.makedirs("static", exist_ok=True)

class AnalyzerManager:
    def __init__(self):
        self.message_queues = {}
        self.active_websockets = {}
        self.input_events = {}
        self.input_responses = {}
        self.lock = threading.Lock()

    async def send_message(self, websocket_id: str, message: dict):
        if websocket_id in self.active_websockets:
            try:
                await self.active_websockets[websocket_id].send_json(message)
            except Exception as e:
                logger.error(f"Error sending message: {e}")

    def register_websocket(self, websocket_id: str, websocket: WebSocket):
        with self.lock:
            self.message_queues[websocket_id] = []
            self.active_websockets[websocket_id] = websocket
            self.input_events[websocket_id] = threading.Event()
            self.input_responses[websocket_id] = None

    def unregister_websocket(self, websocket_id: str):
        with self.lock:
            self.message_queues.pop(websocket_id, None)
            self.active_websockets.pop(websocket_id, None)
            if websocket_id in self.input_events:
                self.input_events[websocket_id].set()
            self.input_events.pop(websocket_id, None)
            self.input_responses.pop(websocket_id, None)

    def get_queue(self, websocket_id: str) -> list:
        return self.message_queues.get(websocket_id)

    def wait_for_input(self, websocket_id: str, timeout=None) -> str:
        if self.input_events[websocket_id].wait(timeout):
            response = self.input_responses[websocket_id]
            self.input_responses[websocket_id] = None
            self.input_events[websocket_id].clear()
            return response
        return None

    def set_input_response(self, websocket_id: str, response: str):
        with self.lock:
            self.input_responses[websocket_id] = response
            self.input_events[websocket_id].set()

analyzer_manager = AnalyzerManager()

@app.get("/")
async def read_root():
    return FileResponse("static/index.html")

@app.get("/download")
async def download_requirements():
    try:
        if not os.path.exists("requirements_answers.txt"):
            raise HTTPException(status_code=404, detail="File not found")
        return FileResponse(
            path="requirements_answers.txt",
            filename="requirements_answers.txt",
            media_type="text/plain"
        )
    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/check_pandoc")
async def check_pandoc():
    """Check if Pandoc is installed and provide installation instructions if not"""
    try:
        # Check if pandoc is installed
        if os.name == 'nt':  # Windows
            process = subprocess.run(["where", "pandoc"], 
                                    stdout=subprocess.PIPE, 
                                    stderr=subprocess.PIPE, 
                                    text=True)
        else:  # Unix/Linux
            process = subprocess.run(["which", "pandoc"], 
                                   stdout=subprocess.PIPE, 
                                   stderr=subprocess.PIPE, 
                                   text=True)
        
        if process.returncode == 0:
            # Pandoc is installed, get version
            version_process = subprocess.run(["pandoc", "--version"], 
                                           stdout=subprocess.PIPE, 
                                           text=True)
            version_info = version_process.stdout.split('\n')[0] if version_process.stdout else "Unknown version"
            
            return JSONResponse({
                "success": True,
                "installed": True,
                "version": version_info,
                "path": process.stdout.strip() if process.stdout else "Unknown path"
            })
        else:
            # Pandoc is not installed
            return JSONResponse({
                "success": True,
                "installed": False,
                "installation_instructions": {
                    "windows": "Download and install from https://pandoc.org/installing.html",
                    "macos": "Run 'brew install pandoc' if you have Homebrew installed",
                    "linux": "Run 'sudo apt-get install pandoc' on Debian/Ubuntu or 'sudo yum install pandoc' on CentOS/RHEL"
                }
            })
    except Exception as e:
        logger.error(f"Error checking Pandoc: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": f"Error checking Pandoc installation: {str(e)}"
        }, status_code=500)

@app.get("/convert_srs")
async def convert_srs(format: str):
    """Convert SRS to different formats"""
    try:
        # Check if the system_srs.md file exists
        if not os.path.exists("system_srs.md"):
            return JSONResponse({
                "success": False,
                "error": "SRS file not found. Please generate SRS document first."
            }, status_code=404)
        
        # Check if pandoc is installed
        try:
            # Use the 'which' command on Unix/Linux or 'where' on Windows to check if pandoc is in PATH
            if os.name == 'nt':  # Windows
                subprocess.run(["where", "pandoc"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            else:  # Unix/Linux
                subprocess.run(["which", "pandoc"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError:
            return JSONResponse({
                "success": False,
                "error": "Pandoc is not installed. Please install Pandoc to convert documents."
            }, status_code=500)
        
        # Convert based on the requested format
        output_file = ""
        if format == "word":
            output_file = "system_srs.docx"
            subprocess.run(["pandoc", "system_srs.md", "-o", output_file], 
                           check=True, stderr=subprocess.PIPE)
        elif format == "pdf":
            output_file = "system_srs.pdf"
            try:
                subprocess.run(["pandoc", "system_srs.md", "-o", output_file], 
                               check=True, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
            except subprocess.CalledProcessError as e:
                error_output = e.stderr.decode('utf-8') if e.stderr else ""
                
                # Check if the error is due to missing pdflatex
                if "pdflatex not found" in error_output:
                    return JSONResponse({
                        "success": False,
                        "error": "PDF generation requires additional software. Please install texlive with: sudo apt-get install texlive-latex-base texlive-fonts-recommended"
                    }, status_code=500)
                else:
                    # Re-raise for the general exception handler
                    raise e
        else:
            return JSONResponse({
                "success": False,
                "error": "Invalid format. Supported formats are 'word' and 'pdf'."
            }, status_code=400)
        
        # Verify the output file was created
        if not os.path.exists(output_file):
            return JSONResponse({
                "success": False,
                "error": f"Conversion failed. Output file '{output_file}' was not created."
            }, status_code=500)
            
        return JSONResponse({"success": True})
    except subprocess.CalledProcessError as e:
        error_message = e.stderr.decode('utf-8') if e.stderr else str(e)
        logger.error(f"Error during Pandoc conversion: {error_message}")
        return JSONResponse({
            "success": False,
            "error": f"Pandoc conversion error: {error_message}"
        }, status_code=500)
    except Exception as e:
        logger.error(f"Error converting SRS: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": f"Error converting document: {str(e)}"
        }, status_code=500)

@app.get("/download_srs")
async def download_srs(format: str = None):
    """Download SRS in various formats"""
    try:
        # Default to markdown if no format is specified
        if not format or format == "md":
            if not os.path.exists("system_srs.md"):
                return JSONResponse({
                    "success": False,
                    "error": "SRS file not found. Please generate SRS document first."
                }, status_code=404)
            
            return FileResponse(
                path="system_srs.md",
                filename="system_srs.md",
                media_type="text/markdown"
            )
        
        # Handle Word format
        elif format == "word":
            if not os.path.exists("system_srs.docx"):
                # Check if Markdown file exists to recommend conversion
                if os.path.exists("system_srs.md"):
                    return JSONResponse({
                        "success": False,
                        "error": "Word file not found. Please convert to Word format first."
                    }, status_code=404)
                else:
                    return JSONResponse({
                        "success": False,
                        "error": "SRS document not found. Please generate SRS document first."
                    }, status_code=404)
            
            return FileResponse(
                path="system_srs.docx",
                filename="system_srs.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        # Handle PDF format
        elif format == "pdf":
            if not os.path.exists("system_srs.pdf"):
                # Check if Markdown file exists to recommend conversion
                if os.path.exists("system_srs.md"):
                    return JSONResponse({
                        "success": False,
                        "error": "PDF file not found. Please convert to PDF format first."
                    }, status_code=404)
                else:
                    return JSONResponse({
                        "success": False,
                        "error": "SRS document not found. Please generate SRS document first."
                    }, status_code=404)
            
            return FileResponse(
                path="system_srs.pdf",
                filename="system_srs.pdf",
                media_type="application/pdf"
            )
        
        # Handle unknown format
        else:
            return JSONResponse({
                "success": False,
                "error": f"Unsupported format: {format}. Supported formats are 'md', 'word', and 'pdf'."
            }, status_code=400)
            
    except Exception as e:
        logger.error(f"Download SRS error: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": f"Error downloading SRS document: {str(e)}"
        }, status_code=500)

@app.post("/create_jira_stories")
async def create_jira_stories():
    """Create Jira stories from SRS"""
    try:
        # Check if system_srs.md exists
        if not os.path.exists("system_srs.md"):
            raise HTTPException(status_code=400, detail="No SRS file found. Please generate SRS first.")
        
        # Read the SRS file
        with open("system_srs.md", "r") as f:
            srs_content = f.read()
        
        # Save as txt for Jira processing
        txt_path = "system_srs.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(srs_content)
        
        # Create a StringIO to capture output
        from io import StringIO
        import sys
        output_capture = StringIO()
        original_stdout = sys.stdout
        sys.stdout = output_capture
        
        try:
            # Import main.py functions
            from main import process_srs_file
            
            # Process the file
            process_srs_file(txt_path)
            
            # Get the captured output
            output = output_capture.getvalue()
            print(f"Processing output: {output}")  # Debug line
        finally:
            sys.stdout = original_stdout
            # Clean up the temporary txt file
            if os.path.exists(txt_path):
                os.remove(txt_path)
        
        return JSONResponse({
            "success": True,
            "output": output
        })
    except Exception as e:
        logger.error(f"Error in Jira integration: {e}")
        return JSONResponse({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        })

@app.get("/download_excel")
async def download_excel():
    """Download the generated Excel file from Jira requirements"""
    try:
        if not os.path.exists("requirements.xlsx"):
            raise HTTPException(status_code=404, detail="Excel file not found")
        return FileResponse("requirements.xlsx", filename="requirements.xlsx")
    except Exception as e:
        logger.error(f"Download Excel error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

def generate_word_document(requirements: list) -> BytesIO:
    """Generate a formatted Word document from requirements"""
    doc = Document()
    
    title = doc.add_heading('Software Requirements Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Requirements: {len(requirements)}')
    
    doc.add_paragraph('Table of Contents', style='Heading 1')
    doc.add_paragraph('(Update table of contents after opening the document)')
    
    req_types = {'FR': 'Functional Requirements',
                'NFR': 'Non-Functional Requirements',
                'SR': 'Security Requirements'}
    
    for req_type, section_title in req_types.items():
        type_reqs = [r for r in requirements if r['type'].startswith(req_type)]
        if type_reqs:
            doc.add_heading(section_title, level=1)
            
            for req in type_reqs:
                heading = doc.add_heading(level=2)
                heading.add_run(f"{req['summary']}")
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.autofit = True
                
                table.columns[0].width = Inches(2)
                table.columns[1].width = Inches(4)
                
                row_cells = table.rows[0].cells
                row_cells[0].text = 'Requirement ID'
                row_cells[1].text = req.get('id', 'N/A')
                
                details = [
                    ('Priority', req.get('priority', 'N/A')),
                    ('Type', req.get('type', 'N/A')),
                    ('Description', req.get('description', 'N/A')),
                ]
                
                for label, value in details:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(value)
                
                doc.add_paragraph()  
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

@app.get("/download_requirement_docx")
async def download_requirement_docx():
    """Download requirements as a formatted Word document"""
    try:
        # Check if requirements_cache.json exists
        if not os.path.exists("requirements_cache.json"):
            raise HTTPException(status_code=404, detail="Requirements cache not found. Please analyze requirements first.")
        
        # Load requirements from cache
        with open("requirements_cache.json", "r") as f:
            requirements_data = json.load(f)
            
        if not requirements_data or not isinstance(requirements_data, list) or len(requirements_data) == 0:
            raise HTTPException(status_code=404, detail="No requirements found in cache.")
            
        # Generate the Word document
        doc_io = generate_word_document(requirements_data)
        
        # Save it temporarily to disk
        with open("requirements_document.docx", "wb") as f:
            f.write(doc_io.getvalue())
        
        return FileResponse(
            path="requirements_document.docx",
            filename="requirements_document.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logger.error(f"Error generating Word document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate_srs_docx")
async def generate_srs_docx(request: Request):
    """Generate a formatted Word document from raw SRS text"""
    try:
        # Get the raw text from the request body
        body = await request.json()
        if not body.get("text"):
            raise HTTPException(status_code=400, detail="No text provided in the request")
        
        raw_text = body.get("text")
        
        # Generate a Word document
        doc = Document()
        
        # Set document style and formatting
        styles = doc.styles
        style_normal = styles['Normal']
        style_normal.font.name = 'Calibri'
        style_normal.font.size = 210000  # 10.5 points
        
        # Add title with center alignment
        title = doc.add_heading('Software Requirements Specification', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process the raw text line by line for better control
        lines = raw_text.split('\n')
        
        # Extract metadata from the first few lines (title, version, date)
        metadata_text = []
        current_line = 0
        
        # Process metadata until we hit a section header
        while current_line < len(lines) and not lines[current_line].startswith('##'):
            line = lines[current_line].strip()
            if line and not line.startswith('#'):
                paragraph = doc.add_paragraph(line)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            current_line += 1
        
        # Add a table of contents
        doc.add_heading('Table of Contents', level=1)
        toc_para = doc.add_paragraph()
        toc_para.add_run('Update table of contents after opening the document')
        
        # Process each section
        section_level = 0
        in_list = False
        list_indent_level = 0
        
        # Continue processing from where we left off
        while current_line < len(lines):
            line = lines[current_line].strip()
            current_line += 1
            
            if not line:
                continue
                
            # Process headings
            if line.startswith('## '):
                # Section header (## 1. Purpose)
                section_name = line[3:].strip()
                heading = doc.add_heading(section_name, level=1)
                in_list = False
            elif line.startswith('### '):
                # Subsection header
                subsection_name = line[4:].strip()
                heading = doc.add_heading(subsection_name, level=2)
                in_list = False
            # Process list items
            elif line.startswith('*'):
                # First level bullet
                list_text = line[1:].strip()
                p = doc.add_paragraph(list_text, style='List Bullet')
                in_list = True
                list_indent_level = 1
            elif line.startswith('    *') or line.startswith('  *'):
                # Indented bullet
                list_text = line.lstrip(' *').strip()
                p = doc.add_paragraph(list_text, style='List Bullet 2')
                in_list = True
                list_indent_level = 2
            # Process bracketed requirement IDs like [FR-001]
            elif '[' in line and ']' in line and any(req_type in line for req_type in ['FR-', 'NFR-', 'SR-']):
                # This appears to be a requirement line
                req_parts = line.split(']', 1)
                if len(req_parts) > 1:
                    req_id = req_parts[0].strip() + ']'
                    req_desc = req_parts[1].strip()
                    
                    # Add a small table for the requirement
                    table = doc.add_table(rows=1, cols=2)
                    table.style = 'Table Grid'
                    
                    # Set column widths
                    table.columns[0].width = Inches(1.5)
                    table.columns[1].width = Inches(5)
                    
                    # Fill the table
                    row_cells = table.rows[0].cells
                    row_cells[0].text = req_id
                    row_cells[1].text = req_desc
                    
                    in_list = False
            else:
                # Regular paragraph
                if in_list and line.startswith(('    ', '  ')):
                    # This is indented text within a list item
                    p = doc.add_paragraph(line.strip())
                    p.style = 'List Continue'
                    p.paragraph_format.left_indent = 30 * list_indent_level * 12000  # Indent based on level
                else:
                    # Regular paragraph
                    doc.add_paragraph(line)
                    in_list = False
        
        # Save the document to a BytesIO object
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        # Save temporarily to disk
        with open("generated_srs.docx", "wb") as f:
            f.write(doc_io.getvalue())
        
        return FileResponse(
            path="generated_srs.docx",
            filename="srs_document.docx",
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        logger.error(f"Error generating SRS document: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

async def process_analyzer_queue(websocket_id: str):
    queue = analyzer_manager.get_queue(websocket_id)
    while True:
        try:
            message = queue.pop(0)
            await analyzer_manager.send_message(websocket_id, message)
        except asyncio.CancelledError:
            break
        except Exception:
            await asyncio.sleep(0.1)

def run_analyzer_in_thread(analyzer: RequirementsAnalyzer, project_description: str, websocket_id: str):
    queue = analyzer_manager.get_queue(websocket_id)
    
    def custom_print(*args, **kwargs):
        message = " ".join(str(arg) for arg in args)
        queue.append({
            "type": "output",
            "message": message
        })

    def custom_input(*args, **kwargs):
        prompt = args[0] if args else "Your answer:"
        queue.append({
            "type": "question",
            "message": prompt
        })
        
        response = analyzer_manager.wait_for_input(websocket_id)
        return response

    try:
        # Replace standard input/output
        import builtins
        original_input = builtins.input
        original_print = builtins.print
        builtins.input = custom_input
        builtins.print = custom_print

        # Run analysis
        results = analyzer.analyze_requirements(project_description)
        
        # Save results to file
        try:
            # First clear the existing file
            with open("requirements_answers.txt", 'w', encoding='utf-8') as f:
                f.write("")  # Clear file
                
            # Then save new results
            analyzer.save_requirements("requirements_answers.txt", results)
            
            queue.append({
                "type": "output",
                "message": "\nRequirements have been saved to 'requirements_answers.txt'"
            })
            queue.append({
                "type": "file_ready",
                "message": "requirements_answers.txt"
            })
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            queue.append({
                "type": "error",
                "message": f"Error saving file: {str(e)}"
            })
            
    except Exception as e:
        logger.error(f"Error in analyzer thread: {str(e)}")
        queue.append({
            "type": "error",
            "message": f"Error during analysis: {str(e)}"
        })
    finally:
        # Restore original input/output
        builtins.input = original_input
        builtins.print = original_print

def run_srs_generator_in_thread(websocket_id: str):
    queue = analyzer_manager.get_queue(websocket_id)
    
    try:
        # Check if requirements file exists and is not empty
        if not os.path.exists("requirements_answers.txt") or os.path.getsize("requirements_answers.txt") == 0:
            queue.append({
                "type": "error",
                "message": "No requirements analysis results found. Please analyze requirements first."
            })
            return
            
        # Run srs_generator_v2.py
        process = subprocess.Popen(
            ["python", "srs_generator_v2.py"],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            bufsize=1,
            universal_newlines=True
        )

        # Stream output in real-time
        while True:
            output = process.stdout.readline()
            if output:
                queue.append({
                    "type": "output",
                    "message": output.strip()
                })
            
            error = process.stderr.readline()
            if error:
                queue.append({
                    "type": "output",
                    "message": "Error: " + error.strip()
                })
            
            if output == '' and error == '' and process.poll() is not None:
                break

        if process.returncode == 0:
            # Check if SRS file was generated
            if os.path.exists("system_srs.md"):
                queue.append({
                    "type": "output",
                    "message": "\nSRS document has been generated successfully!"
                })
                queue.append({
                    "type": "file_ready",
                    "message": "system_srs.md"
                })
            else:
                queue.append({
                    "type": "error",
                    "message": "SRS file was not generated"
                })
        else:
            queue.append({
                "type": "error",
                "message": "SRS generation failed"
            })

    except Exception as e:
        logger.error(f"Error in SRS generator thread: {str(e)}")
        queue.append({
            "type": "error",
            "message": f"Error during SRS generation: {str(e)}"
        })

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    websocket_id = str(id(websocket))
    
    analyzer = RequirementsAnalyzer()
    analyzer_manager.register_websocket(websocket_id, websocket)
    
    queue_processor = asyncio.create_task(process_analyzer_queue(websocket_id))
    
    try:
        while True:
            data = await websocket.receive_json()
            
            if data["type"] == "analyze":
                thread = threading.Thread(
                    target=run_analyzer_in_thread,
                    args=(analyzer, data["description"], websocket_id)
                )
                thread.start()
            elif data["type"] == "generate_srs":
                thread = threading.Thread(
                    target=run_srs_generator_in_thread,
                    args=(websocket_id,)
                )
                thread.start()
            elif data["type"] == "input_response":
                analyzer_manager.set_input_response(websocket_id, data["value"])

    except WebSocketDisconnect:
        logger.info(f"WebSocket disconnected: {websocket_id}")
    except Exception as e:
        logger.error(f"WebSocket error: {str(e)}")
    finally:
        analyzer_manager.unregister_websocket(websocket_id)
        queue_processor.cancel()
        try:
            await websocket.close()
        except:
            pass

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8002)
