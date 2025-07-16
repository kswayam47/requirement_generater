from fastapi import FastAPI, UploadFile, HTTPException, WebSocket, WebSocketDisconnect, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field, validator
import requests
import os
import aiohttp
from dotenv import load_dotenv
import uuid
import google.generativeai as genai  
import json
import traceback
from fastapi.responses import JSONResponse, StreamingResponse
from datetime import datetime, timedelta
from typing import Optional, Dict, List, Set
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from io import BytesIO
import zipfile
import re
import logging
import time
import asyncio
import threading
import subprocess

# Import the RequirementsAnalyzer and NLP
try:
    from requirement_analyzer import RequirementsAnalyzer
except ImportError:
    # Define a simple stub if import fails
    class RequirementsAnalyzer:
        def analyze_requirements(self, *args, **kwargs):
            return "Requirements analysis not available"

# Import the NLP module for requirements analysis
try:
    from nlp import RequirementsGenerator
except ImportError:
    # Define a simple stub if import fails
    class RequirementsGenerator:
        def __init__(self, api_key=None, cache_dir=".cache"):
            pass
        
        def generate_requirements_statement(self, text, format_type="standard"):
            return "NLP analysis not available. Please ensure nlp.py is in the project directory."

logger = logging.getLogger(__name__)
load_dotenv()

app = FastAPI()

# Initialize RequirementsGenerator with API key
nlp_analyzer = RequirementsGenerator(api_key=os.getenv("GEMINI_API_KEY"))

# Mount static files directory
app.mount("/static", StaticFiles(directory="static"), name="static")

# Define BaseModel classes
class JiraIssue(BaseModel):
    summary: str
    description: str
    issue_type: str = Field(default="Story", pattern=r"^(Story|Task|Bug)$")
    priority: str = Field(default="Medium", pattern=r"^(High|Medium|Low)$")

    @validator('priority', pre=True)
    def set_default_priority(cls, value):
        return value or "Medium"  

class EpicRequest(BaseModel):
    epic_summary: str
    epic_description: str
    story_keys: list[str] = []  

class PriorityRequest(BaseModel):
    strategy: str = "business_value"  

class TimelineRequest(BaseModel):
    start_date: datetime
    sprint_duration: int = Field(default=14, description="Sprint duration in days")
    team_capacity: int = Field(default=80, description="Team capacity in story points per sprint")

class DocumentRequest(BaseModel):
    format: str = Field(default="word", pattern="^(word|excel|both)$")
    include_metadata: bool = Field(default=True)
    template_type: str = Field(default="detailed", pattern="^(detailed|simple)$")

# Request models for OCR functionality
class ContentRequest(BaseModel):
    content: str
    type: str = "text"

class UrlRequest(BaseModel):
    url: str

# New model for NLP analysis
class NLPRequest(BaseModel):
    text: str
    format_type: str = "standard"

# WebSocket manager
class AnalyzerManager:
    def __init__(self):
        self.message_queues = {}
        self.active_websockets = {}
        self.input_events = {}
        self.input_responses = {}
        self.lock = threading.Lock()

    async def send_message(self, websocket_id: str, message: dict):
        if websocket_id in self.active_websockets:
            try:
                await self.active_websockets[websocket_id].send_json(message)
            except Exception as e:
                logger.error(f"Error sending message: {e}")

    def register_websocket(self, websocket_id: str, websocket: WebSocket):
        with self.lock:
            self.message_queues[websocket_id] = []
            self.active_websockets[websocket_id] = websocket
            self.input_events[websocket_id] = threading.Event()
            self.input_responses[websocket_id] = None

    def unregister_websocket(self, websocket_id: str):
        with self.lock:
            self.message_queues.pop(websocket_id, None)
            self.active_websockets.pop(websocket_id, None)
            if websocket_id in self.input_events:
                self.input_events[websocket_id].set()
            self.input_events.pop(websocket_id, None)
            self.input_responses.pop(websocket_id, None)

    def get_queue(self, websocket_id: str) -> list:
        return self.message_queues.get(websocket_id, [])

    def wait_for_input(self, websocket_id: str, timeout=None) -> str:
        if self.input_events[websocket_id].wait(timeout):
            response = self.input_responses[websocket_id]
            self.input_responses[websocket_id] = None
            self.input_events[websocket_id].clear()
            return response
        return None

    def set_input_response(self, websocket_id: str, response: str):
        with self.lock:
            self.input_responses[websocket_id] = response
            self.input_events[websocket_id].set()

analyzer_manager = AnalyzerManager()

# WebSocket helper functions
async def process_analyzer_queue(websocket_id: str):
    queue = analyzer_manager.get_queue(websocket_id)
    while True:
        try:
            if queue and len(queue) > 0:
                message = queue.pop(0)
                await analyzer_manager.send_message(websocket_id, message)
        except asyncio.CancelledError:
            break
        except Exception as e:
            logger.error(f"Error processing queue: {e}")
        await asyncio.sleep(0.1)

def run_analyzer_in_thread(analyzer, description, websocket_id):
    try:
        # The actual method only takes the description parameter
        # Let's implement a similar approach to what's in app.py
        message_queue = analyzer_manager.get_queue(websocket_id)
        
        def custom_print(*args, **kwargs):
            message = " ".join(str(arg) for arg in args)
            message_queue.append({
                "type": "output",
                "message": message
            })

        def custom_input(*args, **kwargs):
            prompt = args[0] if args else "Your answer:"
            message_queue.append({
                "type": "question",
                "message": prompt
            })
            
            response = analyzer_manager.wait_for_input(websocket_id)
            return response

        # Replace standard input/output
        import builtins
        original_input = builtins.input
        original_print = builtins.print
        builtins.input = custom_input
        builtins.print = custom_print

        # Run analysis with only the description parameter
        results = analyzer.analyze_requirements(description)
        
        # Save results to file
        try:
            # First clear the existing file
            with open("requirements_answers.txt", 'w', encoding='utf-8') as f:
                f.write("")  # Clear file
                
            # Then save new results if the method exists
            if hasattr(analyzer, 'save_requirements'):
                analyzer.save_requirements("requirements_answers.txt", results)
            else:
                # Fallback if the method doesn't exist
                with open("requirements_answers.txt", 'w', encoding='utf-8') as f:
                    f.write(str(results))
            
            message_queue.append({
                "type": "output",
                "message": "\nRequirements have been saved to 'requirements_answers.txt'"
            })
            message_queue.append({
                "type": "file_ready",
                "message": "requirements_answers.txt"
            })
        except Exception as e:
            logger.error(f"Error saving file: {str(e)}")
            message_queue.append({
                "type": "error",
                "message": f"Error saving file: {str(e)}"
            })
            
    except Exception as e:
        error_msg = f"Error during analysis: {str(e)}"
        logger.error(error_msg)
        message_queue = analyzer_manager.get_queue(websocket_id)
        if message_queue is not None:
            message_queue.append({"type": "error", "message": error_msg})
    finally:
        # Restore original input/output if we modified them
        if 'original_input' in locals() and 'original_print' in locals():
            builtins.input = original_input
            builtins.print = original_print

def run_srs_generator_in_thread(websocket_id):
    try:
        message_queue = analyzer_manager.get_queue(websocket_id)
        
        # Check if requirements file exists and is not empty
        if not os.path.exists("requirements_answers.txt") or os.path.getsize("requirements_answers.txt") == 0:
            message_queue.append({
                "type": "error",
                "message": "No requirements analysis results found. Please analyze requirements first."
            })
            return
            
        # Run srs_generator_v2.py if it exists
        if os.path.exists("srs_generator_v2.py"):
            message_queue.append({
                "type": "output",
                "message": "Generating SRS document..."
            })
            
            try:
                # Run srs_generator_v2.py
                process = subprocess.Popen(
                    ["python", "srs_generator_v2.py"],
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    bufsize=1,
                    universal_newlines=True
                )

                # Stream output in real-time
                while True:
                    output = process.stdout.readline()
                    if output:
                        message_queue.append({
                            "type": "output",
                            "message": output.strip()
                        })
                    
                    error = process.stderr.readline()
                    if error:
                        message_queue.append({
                            "type": "output",
                            "message": "Error: " + error.strip()
                        })
                    
                    if output == '' and error == '' and process.poll() is not None:
                        break

                if process.returncode == 0:
                    # Check if SRS file was generated
                    if os.path.exists("system_srs.md"):
                        message_queue.append({
                            "type": "output",
                            "message": "\nSRS document has been generated successfully!"
                        })
                        message_queue.append({
                            "type": "file_ready",
                            "message": "system_srs.md"
                        })
                    else:
                        message_queue.append({
                            "type": "error",
                            "message": "SRS file was not generated"
                        })
                else:
                    message_queue.append({
                        "type": "error",
                        "message": "SRS generation failed"
                    })
            except Exception as e:
                message_queue.append({
                    "type": "error",
                    "message": f"Error running SRS generator: {str(e)}"
                })
        else:
            # Fallback if srs_generator_v2.py doesn't exist
            message_queue.append({
                "type": "error",
                "message": "SRS generator script not found. Please ensure srs_generator_v2.py exists."
            })
    except Exception as e:
        error_msg = f"Error during SRS generation: {str(e)}"
        logger.error(error_msg)
        message_queue = analyzer_manager.get_queue(websocket_id)
        if message_queue is not None:
            message_queue.append({"type": "error", "message": error_msg})

@app.get("/")
async def read_root():
    return FileResponse("static/index.html")

@app.get("/download")
async def download_requirements():
    try:
        if not os.path.exists("requirements_answers.txt"):
            raise HTTPException(status_code=404, detail="File not found")
        return FileResponse(
            path="requirements_answers.txt",
            filename="requirements_answers.txt",
            media_type="text/plain"
        )
    except Exception as e:
        logger.error(f"Download error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/convert_srs")
async def convert_srs(format: str):
    """Convert SRS to different formats"""
    try:
        # Check if the system_srs.md file exists
        if not os.path.exists("system_srs.md"):
            return JSONResponse({
                "success": False,
                "error": "SRS file not found. Please generate SRS document first."
            }, status_code=404)
        
        # Check if pandoc is installed
        try:
            # Use the 'which' command on Unix/Linux or 'where' on Windows to check if pandoc is in PATH
            if os.name == 'nt':  # Windows
                subprocess.run(["where", "pandoc"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            else:  # Unix/Linux
                subprocess.run(["which", "pandoc"], check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError:
            return JSONResponse({
                "success": False,
                "error": "Pandoc is not installed. Please install Pandoc to convert documents."
            }, status_code=500)
        
        # Convert based on the requested format
        output_file = ""
        if format == "word":
            output_file = "system_srs.docx"
            subprocess.run(["pandoc", "system_srs.md", "-o", output_file], 
                           check=True, stderr=subprocess.PIPE)
        elif format == "pdf":
            output_file = "system_srs.pdf"
            subprocess.run(["pandoc", "system_srs.md", "-o", output_file], 
                           check=True, stderr=subprocess.PIPE)
        else:
            return JSONResponse({
                "success": False,
                "error": "Invalid format. Supported formats are 'word' and 'pdf'."
            }, status_code=400)
        
        # Verify the output file was created
        if not os.path.exists(output_file):
            return JSONResponse({
                "success": False,
                "error": f"Conversion failed. Output file '{output_file}' was not created."
            }, status_code=500)
            
        return JSONResponse({"success": True})
    except subprocess.CalledProcessError as e:
        error_message = e.stderr.decode('utf-8') if e.stderr else str(e)
        logger.error(f"Error during Pandoc conversion: {error_message}")
        return JSONResponse({
            "success": False,
            "error": f"Pandoc conversion error: {error_message}"
        }, status_code=500)
    except Exception as e:
        logger.error(f"Error converting SRS: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": f"Error converting document: {str(e)}"
        }, status_code=500)

@app.get("/download_srs")
async def download_srs(format: str = None):
    """Download SRS in various formats"""
    try:
        # Default to markdown if no format is specified
        if not format or format == "md":
            if not os.path.exists("system_srs.md"):
                return JSONResponse({
                    "success": False,
                    "error": "SRS file not found. Please generate SRS document first."
                }, status_code=404)
            
            return FileResponse(
                path="system_srs.md",
                filename="system_srs.md",
                media_type="text/markdown"
            )
        
        # Handle Word format
        elif format == "word":
            if not os.path.exists("system_srs.docx"):
                # Check if Markdown file exists to recommend conversion
                if os.path.exists("system_srs.md"):
                    return JSONResponse({
                        "success": False,
                        "error": "Word file not found. Please convert to Word format first."
                    }, status_code=404)
                else:
                    return JSONResponse({
                        "success": False,
                        "error": "SRS document not found. Please generate SRS document first."
                    }, status_code=404)
            
            return FileResponse(
                path="system_srs.docx",
                filename="system_srs.docx",
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        
        # Handle PDF format
        elif format == "pdf":
            if not os.path.exists("system_srs.pdf"):
                # Check if Markdown file exists to recommend conversion
                if os.path.exists("system_srs.md"):
                    return JSONResponse({
                        "success": False,
                        "error": "PDF file not found. Please convert to PDF format first."
                    }, status_code=404)
                else:
                    return JSONResponse({
                        "success": False,
                        "error": "SRS document not found. Please generate SRS document first."
                    }, status_code=404)
            
            return FileResponse(
                path="system_srs.pdf",
                filename="system_srs.pdf",
                media_type="application/pdf"
            )
        
        # Handle unknown format
        else:
            return JSONResponse({
                "success": False,
                "error": f"Unsupported format: {format}. Supported formats are 'md', 'word', and 'pdf'."
            }, status_code=400)
            
    except Exception as e:
        logger.error(f"Download SRS error: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": f"Error downloading SRS document: {str(e)}"
        }, status_code=500)

@app.post("/create_jira_stories")
async def create_jira_stories():
    """Create Jira stories from SRS"""
    try:
        # Check if system_srs.md exists
        if not os.path.exists("system_srs.md"):
            raise HTTPException(status_code=400, detail="No SRS file found. Please generate SRS first.")
        
        # Read the SRS file
        with open("system_srs.md", "r") as f:
            srs_content = f.read()
        
        # Save as txt for Jira processing
        txt_path = "system_srs.txt"
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(srs_content)
        
        # Create a StringIO to capture output
        from io import StringIO
        import sys
        output_capture = StringIO()
        original_stdout = sys.stdout
        sys.stdout = output_capture
        
        try:
            # Process the file
            process_srs_file(txt_path)
            
            # Get the captured output
            output = output_capture.getvalue()
            print(f"Processing output: {output}")  # Debug line
        finally:
            sys.stdout = original_stdout
            # Clean up the temporary txt file
            if os.path.exists(txt_path):
                os.remove(txt_path)
        
        return JSONResponse({
            "success": True,
            "output": output
        })
    except Exception as e:
        logger.error(f"Error in Jira integration: {e}")
        return JSONResponse({
            "success": False,
            "error": str(e),
            "traceback": traceback.format_exc()
        })

@app.get("/download_excel")
async def download_excel():
    """Download the generated Excel file from Jira requirements"""
    try:
        if not os.path.exists("requirements.xlsx"):
            raise HTTPException(status_code=404, detail="Excel file not found")
        return FileResponse("requirements.xlsx", filename="requirements.xlsx")
    except Exception as e:
        logger.error(f"Download Excel error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

# OCR and Text Extraction Endpoints
@app.post("/extract_text")
async def extract_text(file: UploadFile = File(...)):
    """Extract text from uploaded files using OCR if needed"""
    try:
        content = await file.read()
        filename = file.filename.lower()
        text = ""
        
        # Create temp directory if it doesn't exist
        os.makedirs("temp", exist_ok=True)
        
        # Process different file types
        if filename.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff')):
            # Save the image temporarily
            temp_path = f"temp/{str(uuid.uuid4())}{os.path.splitext(filename)[1]}"
            with open(temp_path, "wb") as f:
                f.write(content)
            
            try:
                # Use external OCR libraries if available
                try:
                    import pytesseract
                    from PIL import Image
                    
                    # Perform OCR
                    image = Image.open(temp_path)
                    text = pytesseract.image_to_string(image)
                except ImportError:
                    # Fallback to a simpler approach if pytesseract is not available
                    # In a production environment, you'd want to ensure pytesseract is installed
                    text = f"OCR processing requires pytesseract library. Please install it with: pip install pytesseract\n\nAlternatively, you can use PDF files or text-based documents."
            finally:
                # Clean up the temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        elif filename.endswith('.pdf'):
            # Save the PDF temporarily
            temp_path = f"temp/{str(uuid.uuid4())}.pdf"
            with open(temp_path, "wb") as f:
                f.write(content)
            
            try:
                # Extract text from PDF
                try:
                    import PyPDF2
                    
                    reader = PyPDF2.PdfReader(temp_path)
                    text = ""
                    for page_num in range(len(reader.pages)):
                        text += reader.pages[page_num].extract_text() + "\n\n"
                except ImportError:
                    text = f"PDF processing requires PyPDF2 library. Please install it with: pip install PyPDF2"
            finally:
                # Clean up the temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        elif filename.endswith(('.doc', '.docx')):
            # Save the document temporarily
            temp_path = f"temp/{str(uuid.uuid4())}{os.path.splitext(filename)[1]}"
            with open(temp_path, "wb") as f:
                f.write(content)
            
            try:
                # Extract text from Word document
                doc = Document(temp_path)
                text = "\n\n".join([para.text for para in doc.paragraphs])
            finally:
                # Clean up the temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)
        
        elif filename.endswith('.txt'):
            # Plain text file
            text = content.decode('utf-8', errors='replace')
        
        elif filename.endswith(('.html', '.htm')):
            # HTML file
            html_content = content.decode('utf-8', errors='replace')
            
            try:
                from bs4 import BeautifulSoup
                
                soup = BeautifulSoup(html_content, 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.extract()
                
                # Get text
                text = soup.get_text(separator='\n')
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
            except ImportError:
                # Simple HTML parsing without BeautifulSoup
                import re
                html_content = re.sub(r'<[^>]+>', ' ', html_content)
                text = re.sub(r'\s+', ' ', html_content).strip()
        
        else:
            # Unknown file type, try to decode as text
            try:
                text = content.decode('utf-8', errors='replace')
            except:
                raise HTTPException(status_code=400, detail=f"Unsupported file type: {filename}")
        
        # Clean up the text
        text = clean_extracted_text(text)
        
        # Process text through NLP if available
        analyzed_text = ""
        try:
            if text:
                analyzed_text = nlp_analyzer.generate_requirements_statement(text)
        except Exception as nlp_error:
            logger.error(f"NLP processing error: {str(nlp_error)}")
            analyzed_text = f"Error processing text with NLP: {str(nlp_error)}"
        
        return JSONResponse({
            "success": True,
            "text": text,
            "analyzed_text": analyzed_text,
            "filename": filename
        })
    
    except Exception as e:
        logger.error(f"Extract text error: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": str(e)
        })

@app.post("/process_content")
async def process_content(request: ContentRequest):
    """Process content extracted from clipboard or drag-and-drop"""
    try:
        content = request.content
        content_type = request.type
        
        if content_type == "html":
            try:
                from bs4 import BeautifulSoup
                
                soup = BeautifulSoup(content, 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.extract()
                
                # Get text
                text = soup.get_text(separator='\n')
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
            except ImportError:
                # Simple HTML parsing without BeautifulSoup
                import re
                content = re.sub(r'<[^>]+>', ' ', content)
                text = re.sub(r'\s+', ' ', content).strip()
        else:
            # Plain text
            text = content
        
        # Clean up the text
        text = clean_extracted_text(text)
        
        # Process text through NLP if available
        analyzed_text = ""
        try:
            if text:
                analyzed_text = nlp_analyzer.generate_requirements_statement(text)
                # Format the markdown for proper display
                analyzed_text = format_markdown_output(analyzed_text)
        except Exception as nlp_error:
            logger.error(f"NLP processing error: {str(nlp_error)}")
            analyzed_text = f"Error processing text with NLP: {str(nlp_error)}"
        
        return JSONResponse({
            "success": True,
            "text": text,
            "analyzed_text": analyzed_text
        })
    
    except Exception as e:
        logger.error(f"Process content error: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": str(e)
        })

@app.post("/extract_from_url")
async def extract_from_url(request: UrlRequest):
    """Extract content from a URL"""
    try:
        url = request.url
        if not url:
            raise HTTPException(status_code=400, detail="URL is required")
        
        # Check if URL is valid
        if not url.startswith(('http://', 'https://')):
            url = 'https://' + url
        
        async with aiohttp.ClientSession() as session:
            async with session.get(url) as response:
                if response.status != 200:
                    raise HTTPException(status_code=response.status, detail="Failed to fetch URL")
                
                content_type = response.headers.get('Content-Type', '')
                
                if 'text/html' in content_type:
                    # HTML content
                    html = await response.text()
                    
                    try:
                        from bs4 import BeautifulSoup
                        
                        soup = BeautifulSoup(html, 'html.parser')
                        # Remove script and style elements
                        for script in soup(["script", "style"]):
                            script.extract()
                        
                        # Get main content if possible
                        main_content = soup.find('main') or soup.find('article') or soup.find('div', id='content') or soup.find('div', class_='content') or soup.body
                        
                        if main_content:
                            # Get text from main content
                            text = main_content.get_text(separator='\n')
                        else:
                            # Get text from entire document
                            text = soup.get_text(separator='\n')
                        
                        # Clean up whitespace
                        lines = (line.strip() for line in text.splitlines())
                        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                        text = '\n'.join(chunk for chunk in chunks if chunk)
                    except ImportError:
                        # Simple HTML parsing without BeautifulSoup
                        import re
                        html = re.sub(r'<[^>]+>', ' ', html)
                        text = re.sub(r'\s+', ' ', html).strip()
                
                elif 'application/pdf' in content_type:
                    # PDF content
                    pdf_content = await response.read()
                    
                    # Save the PDF temporarily
                    temp_path = f"temp/{str(uuid.uuid4())}.pdf"
                    os.makedirs("temp", exist_ok=True)
                    
                    with open(temp_path, "wb") as f:
                        f.write(pdf_content)
                    
                    try:
                        # Extract text from PDF
                        try:
                            import PyPDF2
                            
                            reader = PyPDF2.PdfReader(temp_path)
                            text = ""
                            for page_num in range(len(reader.pages)):
                                text += reader.pages[page_num].extract_text() + "\n\n"
                        except ImportError:
                            text = f"PDF processing requires PyPDF2 library. Please install it with: pip install PyPDF2"
                    finally:
                        # Clean up the temp file
                        if os.path.exists(temp_path):
                            os.remove(temp_path)
                
                else:
                    # Try to read as plain text
                    text = await response.text()
        
        # Clean up the text
        text = clean_extracted_text(text)
        
        # Process text through NLP if available
        analyzed_text = ""
        try:
            if text:
                analyzed_text = nlp_analyzer.generate_requirements_statement(text)
        except Exception as nlp_error:
            logger.error(f"NLP processing error: {str(nlp_error)}")
            analyzed_text = f"Error processing text with NLP: {str(nlp_error)}"
        
        return JSONResponse({
            "success": True,
            "text": text,
            "analyzed_text": analyzed_text,
            "url": url
        })
    
    except Exception as e:
        logger.error(f"URL extraction error: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": str(e)
        })

def clean_extracted_text(text: str) -> str:
    """Clean up extracted text for better quality"""
    if not text:
        return ""
    
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # Check for excessive newlines - when most lines have just a single word
    lines = text.splitlines()
    if len(lines) > 10:
        # Check if a high percentage of lines contain just 1-2 words
        single_word_lines = sum(1 for line in lines if len(line.strip().split()) <= 2)
        if single_word_lines > len(lines) * 0.7:  # If over 70% of lines are just 1-2 words
            # Join words with spaces, preserving paragraph breaks (2+ newlines)
            # First mark real paragraph breaks
            text = re.sub(r'\n{3,}', '\n\n<PARAGRAPH_BREAK>\n\n', text)
            # Convert single newlines to spaces
            text = re.sub(r'\n(?!\n)', ' ', text)
            # Restore proper paragraph breaks
            text = text.replace('<PARAGRAPH_BREAK>', '')
            # Clean up excessive whitespace
            text = re.sub(r' {2,}', ' ', text)
    else:
        # Replace multiple newlines with a single newline
        text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Clean excessive whitespace but preserve intentional indentation
    text = re.sub(r'[ \t]+$', '', text, flags=re.MULTILINE)  # Remove trailing whitespace
    text = re.sub(r'^[ \t]+$', '', text, flags=re.MULTILINE)  # Remove lines with only whitespace
    
    # Fix bullet points and lists that might have been broken
    text = re.sub(r'([•\-*]) +', r'\1 ', text)
    
    return text.strip()

@app.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await websocket.accept()
    websocket_id = str(id(websocket))
    
    analyzer = RequirementsAnalyzer()
    analyzer_manager.register_websocket(websocket_id, websocket)
    
    queue_processor = asyncio.create_task(process_analyzer_queue(websocket_id))
    
    try:
        while True:
            data = await websocket.receive_json()
            
            if data["type"] == "analyze":
                thread = threading.Thread(
                    target=run_analyzer_in_thread,
                    args=(analyzer, data["description"], websocket_id)
                )
                thread.start()
            elif data["type"] == "generate_srs":
                thread = threading.Thread(
                    target=run_srs_generator_in_thread,
                    args=(websocket_id,)
                )
                thread.start()
            elif data["type"] == "input_response":
                analyzer_manager.set_input_response(websocket_id, data["value"])

    except WebSocketDisconnect:
        logger.info(f"WebSocket disconnected: {websocket_id}")
    except Exception as e:
        logger.error(f"WebSocket error: {str(e)}")
    finally:
        analyzer_manager.unregister_websocket(websocket_id)
        queue_processor.cancel()
        try:
            await websocket.close()
        except:
            pass

# JIRA Configuration
JIRA_URL = os.getenv("JIRA_URL")
JIRA_USER = os.getenv("JIRA_USER")
JIRA_API_TOKEN = os.getenv("JIRA_API_TOKEN")
PROJECT_KEY = os.getenv("JIRA_PROJECT_KEY")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")

genai.configure(api_key=GEMINI_API_KEY)

def convert_to_adf(text: str) -> dict:
    """Convert plain text with line breaks to ADF"""
    content = []
    for paragraph in text.split('\n'):
        if paragraph.strip():
            content.append({
                "type": "paragraph",
                "content": [
                    {
                        "type": "text",
                        "text": paragraph.strip()
                    }
                ]
            })
    return {
        "version": 1,
        "type": "doc",
        "content": content
    }

def create_jira_issue(issue: JiraIssue):
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }

    payload = {
        "fields": {
            "project": {"key": "ARE"},  
            "summary": issue.summary,
            "description": convert_to_adf(issue.description),
            "issuetype": {"name": issue.issue_type}
        }
    }

    if issue.priority:
        payload["fields"]["priority"] = {"name": issue.priority}

    response = requests.post(
        f"{JIRA_URL}/rest/api/3/issue",
        auth=(JIRA_USER, JIRA_API_TOKEN),
        headers=headers,
        json=payload
    )

    if response.status_code != 201:
        raise HTTPException(status_code=response.status_code, detail=response.text)
    
    return response.json()

def create_epic(epic_data: EpicRequest):
    """Create an epic in JIRA Next-gen project"""
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }

    payload = {
        "fields": {
            "project": {"key": "ARE"},
            "summary": epic_data.epic_summary,
            "description": convert_to_adf(epic_data.epic_description),
            "issuetype": {"name": "Epic"}
        }
    }

    response = requests.post(
        f"{JIRA_URL}/rest/api/3/issue",
        auth=(JIRA_USER, JIRA_API_TOKEN),
        headers=headers,
        json=payload
    )

    if response.status_code != 201:
        raise HTTPException(status_code=response.status_code, detail=response.text)
    
    return response.json()

def link_issues(epic_key: str, story_key: str):
    """Link a story to an epic in a Next-gen project"""
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }

    payload = {
        "fields": {
            "parent": {
                "key": epic_key
            }
        }
    }

    response = requests.put(
        f"{JIRA_URL}/rest/api/3/issue/{story_key}",
        auth=(JIRA_USER, JIRA_API_TOKEN),
        headers=headers,
        json=payload
    )

    if response.status_code not in [200, 204]:
        raise HTTPException(status_code=response.status, detail=response.text)

@app.post("/upload-srs/")
async def upload_srs(file: UploadFile):
    """Process SRS document and create JIRA issues"""
    try:
        content = await file.read()
        print("\n" + "="*40 + " ORIGINAL SRS CONTENT " + "="*40)
        print(content.decode())
        print("="*95 + "\n")
        
        requirements = process_srs(content.decode())
        
        created_issues = []
        for req in requirements:
            try:
                print("\n" + "="*40 + " CREATING ISSUE " + "="*40)
                print("Request payload:", req)
                
                issue = JiraIssue(
                    summary=req["summary"],
                    description=req["description"],
                    priority=req.get("priority", "Medium")
                )
                result = create_jira_issue(issue)
                created_issues.append(result)
                
                print("Created issue:", result)
                print("="*95 + "\n")
                
            except Exception as e:
                print(f"Failed to create issue: {str(e)}")
                continue
                
        return {"status": "success", "created_issues": created_issues}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

def identify_domain(content: str) -> str:
    """Use Gemini to identify the domain from the SRS content"""
    model = genai.GenerativeModel("gemini-1.5-pro")
    domain_prompt = f"""Analyze this software requirements document and identify the specific domain (e.g., Healthcare, Banking, Education, etc.).
    Return ONLY the domain name, nothing else.
    
    Document:
    {content[:2000]}  # First 2000 chars should be enough for domain identification
    """
    try:
        response = model.generate_content(domain_prompt)
        domain = response.text.strip()
        return domain
    except Exception as e:
        logger.error(f"Error identifying domain: {e}")
        return "General Software System"

def process_srs(file_content: str):
    """Process structured SRS document and extract requirements"""
    requirements = []
    
    # First identify the domain
    domain = identify_domain(file_content)
    logger.info(f"Identified domain: {domain}")
    
    sections = file_content.split("##")
    req_pattern = r'\[((?:FR|NFR|SR)-\d+)\]:\s*(.*?)\s*\[(High|Medium|Low)\]'
    
    raw_requirements = []
    for section in sections:
        matches = re.finditer(req_pattern, section, re.MULTILINE)
        for match in matches:
            req_id, description, priority = match.groups()
            raw_requirements.append({
                "id": req_id,
                "description": description,
                "priority": priority,
                "type": req_id.split('-')[0],
                "domain": domain
            })
    
    if not raw_requirements:
        return []

    # Process requirements in smaller batches
    BATCH_SIZE = 3
    for i in range(0, len(raw_requirements), BATCH_SIZE):
        batch = raw_requirements[i:i + BATCH_SIZE]
        
        prompt = f"""As an expert in requirements engineering for {domain} systems, enhance these software requirements with clear, professional descriptions and titles.
        For each requirement, provide:
        1. title: A concise, action-oriented title that clearly states the requirement's purpose (max 70 chars)
        2. description: A well-structured description that includes:
           - Clear objective
           - User benefit/business value
           - Technical context
           - Implementation considerations
           - Success criteria
           - Any constraints or dependencies
        
        Format each requirement in a JSON array with this structure:
        {{
            "id": "original_id",
            "title": "enhanced_title",
            "description": "enhanced_description"
        }}

        Original requirements to enhance:
        {json.dumps(batch, indent=2)}
        """
        
        try:
            model = genai.GenerativeModel('gemini-1.5-pro')
            response = model.generate_content(prompt, timeout=30)  # 30 second timeout
            
            # Parse the response
            enhanced = parse_ai_response(response.text)
            if isinstance(enhanced, list):
                requirements.extend(enhanced)
            else:
                print(f"Error parsing batch {i}-{i+BATCH_SIZE}: Invalid response format")
                # Use original requirements as fallback
                for req in batch:
                    requirements.append({
                        "id": req["id"],
                        "title": f"{req['type']} - {req['description'][:50]}...",
                        "description": req["description"]
                    })
        except Exception as e:
            print(f"Error processing batch {i}-{i+BATCH_SIZE}: {str(e)}")
            # Use original requirements as fallback
            for req in batch:
                requirements.append({
                    "id": req["id"],
                    "title": f"{req['type']} - {req['description'][:50]}...",
                    "description": req["description"]
                })
        
        # Add a small delay between batches
        time.sleep(1)
    
    return requirements

@app.post("/prioritize-backlog")
async def prioritize_backlog(strategy: PriorityRequest):
    """AI-driven backlog prioritization with multiple strategies"""
    try:
        print("\n=== STARTING PRIORITIZATION ===")
        
        print("Verifying JIRA configuration...")
        if not all([JIRA_URL, JIRA_USER, JIRA_API_TOKEN, "ARE"]):
            raise ValueError("Missing JIRA configuration in environment variables")
            
        print(f"Fetching backlog for project ARE...")
        async with aiohttp.ClientSession() as session:
            async with session.get(
                f"{JIRA_URL}/rest/api/3/search",
                auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                params={"jql": "project=ARE"}
            ) as response:
                if response.status != 200:
                    text = await response.text()
                    raise HTTPException(status_code=response.status, detail=text)
                
                data = await response.json()
                issues = data.get("issues", [])
        
        print(f"Found {len(issues)} backlog items")
        
        print("Preparing data for AI analysis...")
        issue_data = [{
            "key": issue["key"],
            "summary": issue["fields"]["summary"],
            "description": issue["fields"].get("description", ""),
            "current_priority": issue["fields"]["priority"].get("name", "Medium")
        } for issue in issues]
        
        prompt = f"""Analyze these JIRA issues and prioritize them based on {strategy.strategy}.
Return ONLY a JSON array with this exact structure:
[
    {{
        "issue_key": "PROJ-123", 
        "new_priority": "High",
        "rationale": "1-2 sentence explanation"
    }},
    ...
]
Do NOT include any other text or formatting. Issues:
{json.dumps(issue_data, indent=2)}"""
        
        print("\n=== AI PROMPT ===")
        print(prompt[:1000] + "..." if len(prompt) > 1000 else prompt)
        
        print("Querying Gemini AI...")
        model = genai.GenerativeModel("gemini-1.5-pro")
        response = model.generate_content(prompt)
        raw_response = response.text.replace('```json', '').replace('```', '').strip()
        
        print("\n=== AI RAW RESPONSE ===")
        print(raw_response[:1000] + "..." if len(raw_response) > 1000 else raw_response)
        
        print("Parsing AI response...")
        try:
            prioritized_list = parse_ai_response(raw_response)
            if not isinstance(prioritized_list, list):
                raise ValueError("AI response is not a list")
        except Exception as e:
            print(f"JSON Parsing Error: {str(e)}")
            raise
        
        valid_priorities = ["Highest", "High", "Medium", "Low", "Lowest"]
        for item in prioritized_list:
            if item["new_priority"] not in valid_priorities:
                raise ValueError(f"Invalid priority {item['new_priority']} for {item['issue_key']}")
        
        print("Updating JIRA priorities...")
        success_count = 0
        for item in prioritized_list:
            try:
                async with aiohttp.ClientSession() as session:
                    async with session.put(
                        f"{JIRA_URL}/rest/api/3/issue/{item['issue_key']}",
                        auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                        json={"fields": {"priority": {"name": item["new_priority"]}}}
                    ) as response:
                        if response.status != 204:
                            text = await response.text()
                            print(f"Failed to update {item['issue_key']}: {text}")
                        else:
                            success_count += 1
                    
            except Exception as e:
                print(f"Error updating {item['issue_key']}: {str(e)}")
        
        return {
            "status": "partial" if success_count < len(prioritized_list) else "success",
            "updated": success_count,
            "failed": len(prioritized_list) - success_count
        }

    except Exception as e:
        print("\n=== ERROR DETAILS ===")
        print(f"Type: {type(e).__name__}")
        print(f"Message: {str(e)}")
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail={
            "error": str(e),
            "type": type(e).__name__
        })

@app.post("/create-epic/")
async def create_epic_endpoint(epic_request: EpicRequest):
    """Create an epic and link stories to it"""
    try:
        epic_result = create_epic(epic_request)
        epic_key = epic_result["key"]
        
        linked_stories = []
        for story_key in epic_request.story_keys:
            try:
                link_issues(epic_key, story_key)
                linked_stories.append(story_key)
            except Exception as e:
                print(f"Failed to link story {story_key}: {str(e)}")
        
        return {
            "status": "success",
            "epic_key": epic_key,
            "linked_stories": linked_stories,
            "failed_stories": [key for key in epic_request.story_keys if key not in linked_stories]
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create epic or link stories: {str(e)}"
        )

@app.post("/automate-workflow/")
async def automate_workflow():
    """Automate task assignment and status updates"""
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(
                f"{JIRA_URL}/rest/api/3/search",
                auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                params={"jql": f"project={PROJECT_KEY} AND status='To Do'"}
            ) as response:
                if response.status != 200:
                    text = await response.text()
                    raise HTTPException(status_code=response.status, detail=text)
                
                data = await response.json()
                issues = data.get("issues", [])
        
        for issue in issues:
            issue_key = issue["key"]
            transition_payload = {"transition": {"id": "21"}}  
            
            async with aiohttp.ClientSession() as session:
                async with session.post(
                    f"{JIRA_URL}/rest/api/3/issue/{issue_key}/transitions",
                    auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                    json=transition_payload
                ) as response:
                    if response.status != 204:
                        text = await response.text()
                        raise HTTPException(status_code=response.status, detail=text)
        
        return {"status": "success", "message": "Workflow automated successfully"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

async def get_jira_issues(jql: str = "issuetype = Story"):
    """Fetch issues from JIRA based on JQL"""
    headers = {
        "Accept": "application/json",
        "Content-Type": "application/json"
    }
    
    params = {
        "jql": jql,
        "maxResults": 100  
    }
    
    async with aiohttp.ClientSession() as session:
        async with session.get(
            f"{JIRA_URL}/rest/api/3/search",
            auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
            headers=headers,
            params=params
        ) as response:
            if response.status != 200:
                text = await response.text()
                raise HTTPException(status_code=response.status, detail=text)
            
            data = await response.json()
            return data.get("issues", [])

def suggest_epics(stories: list) -> list:
    """Use Gemini AI to suggest epic groupings for stories"""
    stories_text = "\n".join([
        f"Story {story['key']}:\n"
        f"Title: {story['fields']['summary']}\n"
        f"Description: {story['fields'].get('description', 'No description')}\n"
        f"Priority: {story['fields'].get('priority', {}).get('name', 'Not set')}\n"
        f"Labels: {', '.join(story['fields'].get('labels', []))}\n"
        for story in stories
    ])
    
    prompt = f"""As an expert in Agile project management and JIRA organization, analyze these user stories and suggest optimal epic groupings.

    Consider these key factors when grouping stories into epics:
    1. Business Value Stream:
       - Group stories that contribute to the same business objective
       - Consider end-user value and customer journey
       - Identify common business processes or workflows
    
    2. Technical Dependencies:
       - Identify stories that share technical components
       - Consider architectural layers (UI, backend, database)
       - Group stories with similar technical requirements
    
    3. Feature Completeness:
       - Ensure each epic represents a complete, deliverable feature
       - Consider MVP (Minimum Viable Product) requirements
       - Include all necessary stories for end-to-end functionality
    
    4. Timeline and Priority:
       - Group stories that should be implemented in the same timeframe
       - Consider dependencies and sequential implementation needs
       - Balance epic sizes for manageable delivery
    
    For each suggested epic, provide:
    1. epic_summary: A clear, concise title (max 70 chars)
    2. epic_description: A detailed description including:
       - Epic's main objective
       - Key deliverables
       - Success criteria
       - Technical considerations
       - Dependencies on other epics
    3. story_keys: List of story keys to include
    4. rationale: Detailed explanation including:
       - Why these stories belong together
       - Business value alignment
       - Technical synergies
       - Implementation considerations
    5. estimated_duration: Rough time estimate (in sprints)
    6. suggested_priority: High/Medium/Low based on business impact
    
    Return the response as a JSON array where each epic follows this structure:
    {{
        "epic_summary": "string",
        "epic_description": "string",
        "story_keys": ["key1", "key2"],
        "rationale": "string",
        "estimated_duration": "number",
        "suggested_priority": "string"
    }}

    Here are the stories to analyze:
    {stories_text}
    """
    
    model = genai.GenerativeModel("gemini-1.5-pro")
    response = model.generate_content(prompt)
    
    try:
        return parse_ai_response(response.text)
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to parse AI suggestions: {str(e)}"
        )

@app.post("/suggest-epics/")
async def suggest_epics_endpoint(jql: str = "issuetype = Story"):
    """Get AI suggestions for grouping stories into epics"""
    try:
        # Fetch stories from JIRA
        stories = await get_jira_issues(jql)
        
        if not stories:
            return {"message": "No stories found matching the criteria"}
        
        # Get AI suggestions
        suggestions = suggest_epics(stories)
        
        return {
            "status": "success",
            "epic_suggestions": suggestions,
            "total_stories_analyzed": len(stories)
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate epic suggestions: {str(e)}"
        )

@app.post("/create-suggested-epics/")
async def create_suggested_epics(jql: str = "issuetype = Story"):
    """Get AI suggestions and automatically create the suggested epics"""
    try:
        suggestions = await suggest_epics_endpoint(jql)
        
        if "epic_suggestions" not in suggestions:
            return suggestions
        
        created_epics = []
        for epic in suggestions["epic_suggestions"]:
            epic_request = EpicRequest(
                epic_summary=epic["epic_summary"],
                epic_description=f"{epic['epic_description']}\n\nRationale: {epic['rationale']}",
                story_keys=epic["story_keys"]
            )
            
            result = await create_epic_endpoint(epic_request)
            created_epics.append(result)
        
        return {
            "status": "success",
            "created_epics": created_epics,
            "total_epics_created": len(created_epics)
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create suggested epics: {str(e)}"
        )

async def generate_epic_timeline(epics: list, timeline_req: TimelineRequest):
    """Generate timeline for epics using GenAI"""
    epics_data = []
    for epic in epics:
        epic_data = {
            "key": epic["key"],
            "summary": epic["fields"]["summary"],
            "description": epic["fields"].get("description", ""),
            "priority": epic["fields"].get("priority", {}).get("name", "Medium"),
            "stories": []
        }
        
        jql = f'parent = {epic["key"]}'
        stories = await get_jira_issues(jql)
        epic_data["stories"] = [{
            "key": story["key"],
            "summary": story["fields"]["summary"],
            "story_points": story["fields"].get("customfield_10014", 0)  
        } for story in stories]
        
        epics_data.append(epic_data)

    prompt = f"""As an expert in Agile project planning, create an optimal timeline for these epics starting from {timeline_req.start_date.strftime('%Y-%m-%d')}.

    Project Context:
    - Sprint duration: {timeline_req.sprint_duration} days
    - Team capacity: {timeline_req.team_capacity} story points per sprint
    
    For each epic, analyze:
    1. Scope and complexity
    2. Dependencies between stories
    3. Priority and business value
    4. Required team capacity
    5. Technical dependencies
    
    Epics and their stories:
    {json.dumps(epics_data, indent=2)}
    
    Return a JSON array of epic schedules with:
    {{
        "epic_key": "string",
        "start_date": "YYYY-MM-DD",
        "end_date": "YYYY-MM-DD",
        "estimated_sprints": number,
        "story_point_total": number,
        "parallel_epics": ["epic_keys"],
        "scheduling_rationale": "string"
    }}
    
    Consider:
    - Story point distribution across sprints
    - Dependencies between epics
    - Parallel execution possibilities
    - Team capacity constraints
    - Priority order
    """
    
    model = genai.GenerativeModel("gemini-1.5-pro")
    response = model.generate_content(prompt)
    
    try:
        timeline = parse_ai_response(response.text)
        
        for schedule in timeline:
            epic_key = schedule["epic_key"]
            
            update_payload = {
                "fields": {
                    "customfield_10015": schedule["start_date"],  
                    "duedate": schedule["end_date"],  
                    "description": {
                        "version": 1,
                        "type": "doc",
                        "content": [
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"Estimated Duration: {schedule['estimated_sprints']} sprints\n"
                                    }
                                ]
                            },
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"Story Points: {schedule['story_point_total']}\n"
                                    }
                                ]
                            },
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"Parallel Epics: {', '.join(schedule['parallel_epics'])}\n"
                                    }
                                ]
                            },
                            {
                                "type": "paragraph",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": f"\nScheduling Rationale:\n{schedule['scheduling_rationale']}"
                                    }
                                ]
                            }
                        ]
                    }
                }
            }
            
            headers = {
                "Accept": "application/json",
                "Content-Type": "application/json"
            }
            
            async with aiohttp.ClientSession() as session:
                async with session.put(
                    f"{JIRA_URL}/rest/api/3/issue/{epic_key}",
                    auth=aiohttp.BasicAuth(JIRA_USER, JIRA_API_TOKEN),
                    headers=headers,
                    json=update_payload
                ) as response:
                    if response.status != 204:
                        text = await response.text()
                        print(f"Failed to update epic {epic_key}: {text}")
        
        return timeline
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate timeline: {str(e)}"
        )

@app.post("/generate-timeline")
async def generate_timeline_endpoint(timeline_req: TimelineRequest):
    """Generate and apply timeline for all epics"""
    try:
        epics = await get_jira_issues("issuetype = Epic")
        
        timeline = await generate_epic_timeline(epics, timeline_req)
        
        return {
            "status": "success",
            "timeline": timeline,
            "start_date": timeline_req.start_date,
            "total_epics": len(timeline)
        }
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate timeline: {str(e)}"
        )

async def generate_word_document(requirements: list) -> BytesIO:
    """Generate a formatted Word document from requirements"""
    doc = Document()
    
    title = doc.add_heading('Software Requirements Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Requirements: {len(requirements)}')
    
    doc.add_paragraph('Table of Contents', style='Heading 1')
    doc.add_paragraph('(Update table of contents after opening the document)')
    
    req_types = {'FR': 'Functional Requirements',
                'NFR': 'Non-Functional Requirements',
                'SR': 'Security Requirements'}
    
    for req_type, section_title in req_types.items():
        type_reqs = [r for r in requirements if r['type'].startswith(req_type)]
        if type_reqs:
            doc.add_heading(section_title, level=1)
            
            for req in type_reqs:
                heading = doc.add_heading(level=2)
                heading.add_run(f"{req['summary']}")
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.autofit = True
                
                table.columns[0].width = Inches(2)
                table.columns[1].width = Inches(4)
                
                row_cells = table.rows[0].cells
                row_cells[0].text = 'Requirement ID'
                row_cells[1].text = req.get('id', 'N/A')
                
                details = [
                    ('Priority', req.get('priority', 'N/A')),
                    ('Type', req.get('type', 'N/A')),
                    ('Description', req.get('description', 'N/A')),
                ]
                
                for label, value in details:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(value)
                
                doc.add_paragraph()  
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

async def parse_description_fields(description: str | dict) -> dict:
    """Parse structured fields from description"""
    fields = {
        'Objective': '',
        'User Benefit': '',
        'Technical Context': '',
        'Implementation Considerations': '',
        'Success Criteria': '',
        'Constraints': ''
    }
    
    if isinstance(description, dict):
        description = convert_adf_to_text(description)
    elif description is None:
        description = ''
    
    if description:
        for field in fields.keys():
            pattern = rf"\*\*{field}:\*\*(.*?)(?=\*\*\w+:\*\*|$)"
            match = re.search(pattern, description, re.DOTALL)
            if match:
                fields[field] = match.group(1).strip()
    
    return fields

def convert_adf_to_text(adf: dict) -> str:
    """Convert Atlassian Document Format to plain text"""
    text = []
    
    if not adf or 'content' not in adf:
        return ''
    
    for item in adf.get('content', []):
        if item.get('type') == 'paragraph':
            paragraph_text = []
            for content in item.get('content', []):
                if content.get('type') == 'text':
                    paragraph_text.append(content.get('text', ''))
            text.append(''.join(paragraph_text))
    
    return '\n'.join(text)

def extract_original_requirement(description: str | dict) -> str:
    """Extract the original requirement value from the description"""
    if isinstance(description, dict):
        description = convert_adf_to_text(description)
    elif description is None:
        return ''
    
    pattern = r'\[((?:FR|NFR|SR)-\d+)\]:\s*(.*?)\s*\[(High|Medium|Low)\]'
    match = re.search(pattern, description)
    if match:
        return match.group(2).strip()
    return description  

async def generate_excel_sheet(requirements: list) -> BytesIO:
    """Generate an Excel sheet from requirements for JIRA import"""

    excel_data = []
    for req in requirements:

        epic_link = ''
        if req['type'] == 'FR':
            epic_link = 'ARE-123'  
        elif req['type'] == 'NFR':
            epic_link = 'ARE-124'  
        elif req['type'] == 'SR':
            epic_link = 'ARE-125' 
        
        original_req = extract_original_requirement(req.get('description', ''))
        
        components = {
            'FR': 'Functional',
            'NFR': 'Non-Functional',
            'SR': 'Security'
        }.get(req['type'], '')
        
        excel_data.append({
            'Summary': req['summary'],
            'Description': original_req,
            'Priority': req['priority'],
            'Components': components,
            'Original Requirement ID': req.get('id', ''),
            'Epic Link': epic_link,
            'Is_Epic': req.get('fields', {}).get('issuetype', {}).get('name') == 'Epic'  
        })
    
    df = pd.DataFrame(excel_data)
    is_epic = df.pop('Is_Epic')  
    
    excel_io = BytesIO()
    with pd.ExcelWriter(excel_io, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='JIRA Import', index=False)
        
        workbook = writer.book
        worksheet = writer.sheets['JIRA Import']
        
        header_format = workbook.add_format({
            'bold': True,
            'bg_color': '#D8E4BC',
            'border': 1
        })
        
        epic_format = workbook.add_format({
            'text_wrap': True,
            'valign': 'top',
            'bg_color': '#F0F3F7'  
        })
        
        story_format = workbook.add_format({
            'text_wrap': True,
            'valign': 'top',
            'bg_color': '#E8F3E8'  
        })
        
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
        
        for row in range(len(df)):
            row_format = epic_format if is_epic.iloc[row] else story_format
            for col in range(len(df.columns)):
                worksheet.write(row + 1, col, df.iloc[row, col], row_format)
            
        worksheet.set_column('A:A', 40)  
        worksheet.set_column('B:B', 60)  
        worksheet.set_column('C:F', 15)  
        
        worksheet.set_default_row(100)
    
    excel_io.seek(0)
    return excel_io

@app.post("/generate-documents")
async def generate_documents(doc_request: DocumentRequest):
    """Generate requirement documents in Word and/or Excel format"""
    try:
        requirements = []
        epics = await get_jira_issues("issuetype = Epic")
        stories = await get_jira_issues("issuetype = Story")
        
        for item in epics + stories:
            req = {
                'id': item['key'],
                'summary': item['fields']['summary'],
                'description': item['fields'].get('description', ''),
                'type': 'FR' if item['fields']['issuetype']['name'] == 'Story' else 'NFR',
                'priority': item['fields'].get('priority', {}).get('name', 'Medium')
            }
            requirements.append(req)
        
        if doc_request.format == 'word':
            word_io = await generate_word_document(requirements)
            return StreamingResponse(
                word_io,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": "attachment; filename=requirements.docx"
                }
            )
            
        elif doc_request.format == 'excel':
            excel_io = await generate_excel_sheet(requirements)
            return StreamingResponse(
                excel_io,
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                headers={
                    "Content-Disposition": "attachment; filename=jira_import.xlsx"
                }
            )
            
        else:  
            zip_io = BytesIO()
            with zipfile.ZipFile(zip_io, mode='w', compression=zipfile.ZIP_DEFLATED) as zip_file:
    
                word_io = await generate_word_document(requirements)
                zip_file.writestr('requirements.docx', word_io.getvalue())

                excel_io = await generate_excel_sheet(requirements)
                zip_file.writestr('jira_import.xlsx', excel_io.getvalue())
            
            zip_io.seek(0)
            return StreamingResponse(
                zip_io,
                media_type="application/zip",
                headers={
                    "Content-Disposition": "attachment; filename=requirement_documents.zip"
                }
            )
        
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to generate documents: {str(e)}"
        )

def parse_ai_response(raw_response: str) -> list:
    """Safely parse Gemini's JSON response with potential markdown formatting"""
    try:
        clean_response = re.sub(r'```json|```', '', raw_response)
        
        start = clean_response.find('[')
        end = clean_response.rfind(']') + 1
        
        if start == -1 or end == 0:
            raise ValueError("No JSON array found in response")
            
        json_str = clean_response[start:end]
        
        parsed = json.loads(json_str)
        
        if not isinstance(parsed, list):
            raise ValueError("Response is not a JSON array")
            
        return parsed
        
    except Exception as e:
        print("Failed to parse AI response:")
        print(f"Raw response: {raw_response}")
        print(f"Clean attempt: {json_str if 'json_str' in locals() else 'N/A'}")
        raise

def process_srs_file(file_path: str):
    """Process SRS file and create Jira stories"""
    try:
        # Read the file content
        with open(file_path, 'r') as f:
            file_content = f.read()
        
        # Process the requirements
        requirements = process_srs(file_content)
        
        print(f"Processing {len(requirements)} requirements...")
        
        # Create Jira issues for each requirement
        for req in requirements:
            try:
                # Map requirement type to issue type
                issue_type_map = {
                    "FR": "Story",
                    "NFR": "Task",
                    "SR": "Task"
                }
                
                # Extract type from ID (e.g., "FR-1" -> "FR")
                req_type = req["id"].split("-")[0]
                
                issue = JiraIssue(
                    summary=f"{req['id']}: {req['title']}",
                    description=(
                        f"*Requirement ID:* {req['id']}\n\n"
                        f"*Type:* {req_type}\n\n"
                        f"*Description:*\n{req['description']}\n\n"
                    ),
                    issue_type=issue_type_map.get(req_type, "Story"),
                    priority="Medium"
                )
                response = create_jira_issue(issue)
                print(f"Created Jira issue: {response['key']} - {req['title']}")
            except Exception as e:
                print(f"Error creating Jira issue for {req['id']}: {str(e)}")
                continue
        
        print("Completed processing requirements.")
        return True
    except Exception as e:
        print(f"Error processing SRS file: {str(e)}")
        raise

@app.post("/analyze_requirements")
async def analyze_requirements(request: NLPRequest):
    """Analyze text using NLP to extract requirements"""
    try:
        if not request.text:
            raise HTTPException(status_code=400, detail="Text is required")
        
        # Process the text with NLP
        analyzed_text = nlp_analyzer.generate_requirements_statement(
            request.text, 
            format_type=request.format_type
        )
        
        return JSONResponse({
            "success": True,
            "original_text": request.text,
            "analyzed_text": analyzed_text
        })
    
    except Exception as e:
        logger.error(f"NLP analysis error: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": str(e)
        })

@app.post("/normalize_text")
async def normalize_text(request: NLPRequest):
    """Process text and normalize formatting to handle excessive line breaks"""
    try:
        text = request.text
        if not text:
            return JSONResponse({
                "success": False,
                "error": "No text provided"
            })
        
        # Apply special cleaning for excessive newlines
        cleaned_text = clean_extracted_text(text)
        
        return JSONResponse({
            "success": True,
            "text": cleaned_text,
        })
    
    except Exception as e:
        logger.error(f"Text normalization error: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": str(e)
        })

@app.get("/check_pandoc")
async def check_pandoc():
    """Check if Pandoc is installed and provide installation instructions if not"""
    try:
        # Check if pandoc is installed
        if os.name == 'nt':  # Windows
            process = subprocess.run(["where", "pandoc"], 
                                    stdout=subprocess.PIPE, 
                                    stderr=subprocess.PIPE, 
                                    text=True)
        else:  # Unix/Linux
            process = subprocess.run(["which", "pandoc"], 
                                   stdout=subprocess.PIPE, 
                                   stderr=subprocess.PIPE, 
                                   text=True)
        
        if process.returncode == 0:
            # Pandoc is installed, get version
            version_process = subprocess.run(["pandoc", "--version"], 
                                           stdout=subprocess.PIPE, 
                                           text=True)
            version_info = version_process.stdout.split('\n')[0] if version_process.stdout else "Unknown version"
            
            return JSONResponse({
                "success": True,
                "installed": True,
                "version": version_info,
                "path": process.stdout.strip() if process.stdout else "Unknown path"
            })
        else:
            # Pandoc is not installed
            return JSONResponse({
                "success": True,
                "installed": False,
                "installation_instructions": {
                    "windows": "Download and install from https://pandoc.org/installing.html",
                    "macos": "Run 'brew install pandoc' if you have Homebrew installed",
                    "linux": "Run 'sudo apt-get install pandoc' on Debian/Ubuntu or 'sudo yum install pandoc' on CentOS/RHEL"
                }
            })
    except Exception as e:
        logger.error(f"Error checking Pandoc: {str(e)}")
        return JSONResponse({
            "success": False,
            "error": f"Error checking Pandoc installation: {str(e)}"
        }, status_code=500)

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)
