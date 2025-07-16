import requests
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime
from io import BytesIO

def process_srs_file(file_path):
    """Process the SRS file and create Jira issues, then export to Excel"""
    
    # Check if file exists
    if not Path(file_path).exists():
        print(f"Error: File {file_path} not found!")
        return
    
    # Upload and process SRS
    print("Uploading and processing SRS file...")
    files = {'file': open(file_path, 'rb')}
    response = requests.post('http://localhost:8000/upload-srs/', files=files)
    
    if response.status_code == 200:
        print("SRS processed successfully!")
        result = response.json()
        print(f"\nCreated {len(result.get('created_issues', []))} issues:")
        for issue in result.get('created_issues', []):
            print(f"- {issue.get('key')}: {issue.get('summary')}")
        
        # Generate Excel document
        print("\nGenerating Excel document...")
        try:
            doc_response = requests.post(
                'http://localhost:8000/generate-documents', 
                json={
                    "format": "excel",
                    "include_metadata": True,
                    "template_type": "simple"
                },
                stream=True  # Stream the response
            )
            
            if doc_response.status_code == 200:
                # Save the file
                filename = 'requirements.xlsx'
                with open(filename, 'wb') as f:
                    for chunk in doc_response.iter_content(chunk_size=8192):
                        if chunk:
                            f.write(chunk)
                print(f"\nExcel file '{filename}' has been created successfully!")
                print(f"You can find it at: {os.path.abspath(filename)}")
            else:
                print(f"Error generating Excel: {doc_response.status_code}")
                print(doc_response.text)
        except Exception as e:
            print(f"Error creating Excel file: {str(e)}")
    else:
        print(f"Error: {response.status_code}")
        print(response.text)

def extract_sections(srs_content):
    """Extract sections from SRS markdown content"""
    # ... existing code ...

def generate_word_document(requirements: list) -> BytesIO:
    """Generate a formatted Word document from requirements"""
    doc = Document()
    
    title = doc.add_heading('Software Requirements Specification', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph(f'Generated Date: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'Total Requirements: {len(requirements)}')
    
    doc.add_paragraph('Table of Contents', style='Heading 1')
    doc.add_paragraph('(Update table of contents after opening the document)')
    
    req_types = {'FR': 'Functional Requirements',
                'NFR': 'Non-Functional Requirements',
                'SR': 'Security Requirements'}
    
    for req_type, section_title in req_types.items():
        type_reqs = [r for r in requirements if r['type'].startswith(req_type)]
        if type_reqs:
            doc.add_heading(section_title, level=1)
            
            for req in type_reqs:
                heading = doc.add_heading(level=2)
                heading.add_run(f"{req['summary']}")
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                table.autofit = True
                
                table.columns[0].width = Inches(2)
                table.columns[1].width = Inches(4)
                
                row_cells = table.rows[0].cells
                row_cells[0].text = 'Requirement ID'
                row_cells[1].text = req.get('id', 'N/A')
                
                details = [
                    ('Priority', req.get('priority', 'N/A')),
                    ('Type', req.get('type', 'N/A')),
                    ('Description', req.get('description', 'N/A')),
                ]
                
                for label, value in details:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    row_cells[1].text = str(value)
                
                doc.add_paragraph()  
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def generate_srs_document(srs_content, output_file="system_srs.docx"):
    """Generate a formatted DOCX document from SRS content"""
    # ... existing code ...

if __name__ == "__main__":
    srs_file = "system_srs.txt"
    process_srs_file(srs_file)
